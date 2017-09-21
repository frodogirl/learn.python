from docx import Document
import html5lib
from html5lib import treebuilders, treewalkers

from referat import get_topic, get_unique_referat


def parse_referat(html):
    """
        TODO: При помощи библиотеки html5lib https://html5lib.readthedocs.io/en/latest/
        разобрать текст реферата на блоки и положить в словарь
        Содержимое тега div - в topic (тема)
        Содержимое тега strong - в title (название)
        Содержимое тегов p - отдельными строками в список content (текст)
    """
    parser = html5lib.HTMLParser(tree=html5lib.getTreeBuilder("dom"))
    dom_tree = parser.parse(html)
    
    topic = dom_tree.getElementsByTagName('div')[0].firstChild.nodeValue
    title = dom_tree.getElementsByTagName('strong')[0].firstChild.nodeValue
    content = []
    p_tags = dom_tree.getElementsByTagName('p')
    for p_tag in p_tags:
        content.append(p_tag.firstChild.nodeValue)

    return {'topic': topic, 'title': title, 'content': content}

def save_referat(referat, filename):
    """
        TODO: сохранить реферат в .docx файл при помощи python-docx https://python-docx.readthedocs.io/en/latest/
        topic вывести на отдельном листе большим шрифтом
        title - как заголовок, размер меньше, чем у topic
        Каждую строку из content - как отдельный абзац
    """
    document = Document()

    document.add_heading(referat.get('topic'))
    document.add_page_break()
    document.add_heading(referat.get('title'))
    content = referat.get('content')
    for paragraph in content:
        document.add_paragraph(paragraph)
    
    document.save(filename)


if __name__ == '__main__':
    topic = get_topic()
    html_data = get_unique_referat(topic)
    referat = parse_referat(html_data)
    save_referat(referat, 'referat.docx')
    print('ok')
